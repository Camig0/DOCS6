from docx import  Document
import tkinter
from tkinter import filedialog
from bs4 import BeautifulSoup
from dataclasses import dataclass

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import RGBColor

import time
import os
#TODO: comments/documentation
#TODO: paragraph styles



class SuperElement: #for now this where all the defaults stay and mayve here should also be the default adding of each run
    #                       - it can be that it is a function that be passed in a function
    def __init__(self):
        self.PRESET_COLORS = {
            'black': RGBColor(0,0,0),
            'red': RGBColor(255,0,0),
            'green': RGBColor(0,255,0),
            'blue': RGBColor(0,0,255),
            'yellow': RGBColor(255,255,0),
            'cyan': RGBColor(0,255,255),
            'magenta': RGBColor(255,0,255),
            'white': RGBColor(255,255,255),
        }
        self.DEFAULT_STYLE:dict =  { #defaults should be set here #FUTURE
            'bold': False, #bool
            'italic': False, #bool
            'underline':False, #bool
            'highlight': None, #int from 1-5;
            'color': 'black', #colors from PRESET_COLORS
            'font_size': 12, #int
        }
    def addRuntoRootContainer(self,root_paragraph,styles,run):
        for style, value in styles.items():
            print(style,value)
            if style == 'bold':
                run.bold = value
            if style == 'italic':
                run.italic = value
            if style == 'underline':
                run.underline = value
            if style == 'font_size':
                run.font.size = value
            if style == 'color':
                run.font.color.rgb = self.PRESET_COLORS.get(value)


@dataclass
class DocList:
    rootParagraph:any
    ...


@dataclass
class DocRun(SuperElement):
    def __init__(self, _rootParagraph, _element):
        self.rootParagraph = _rootParagraph
        self.run = self.rootParagraph.add_run(_element.get('element'))
        self.element:str = _element.get('element')
        super().__init__()
        parent_tags:tuple[dict] = _element.get('parent_tags')

        #1. get styles of the element

        styles:dict =  self.DEFAULT_STYLE

        for parent_tag in parent_tags:
            #split attributes to key to value ex. -> attributes = [('bold', True), ('italic', True), ('underline', True)]
            attributes:list[str] = parent_tag[0].get('attributes') 
            attributes = map(lambda x:x.replace(',',''), attributes) #---> removes the annoying commas if there are any
            attributes = [tuple(v.split('=')) for v in attributes]
#TODO: add check here if attribute tuple have 2 values if only one is full fill with 'None'
            for key, value in attributes:
                if key not in styles:
                    continue #filters the attributes that are not in the styles dictionary
                elif value.lower() == 'true': #--> if elif chain for setting the values appropriately
                    styles[key] = True
                elif value.lower() == 'false':
                    styles[key] = False
                elif value.isdigit():
                    styles[key] = Pt(int(value))
                    ...
                else:
                    styles[key] = value

        self.styles = styles
        return None

        #2. execute styles

    def addtoRootParagraph(self) -> None:
        #using the styles given apply them to a run in a paragraph
            SuperElement.addRuntoRootContainer(self,self.rootParagraph,self.styles,self.run)


#goal of this file is to make a word document from given tags

def get_target_dir() -> str:
    """
    Opens a dialog to select a target directory.

    Returns the path of the selected directory.

    Returns:
    str: The path of the selected directory.11
    """
    tkinter.Tk().withdraw()  # prevents an empty tkinter window from appearing
    folder_path = filedialog.askdirectory()
    return folder_path

def get_tags(text:str)-> list[dict[str, any]]:
    """
    Extracts tags from the given text content and returns a list of dictionaries representing each tag.

    Parameters:
    text (str): The text content containing HTML-like tags.

    Returns:
    list: A list of dictionaries representing tags with keys 'start', 'end', 'structure', 'type', 'attributes', and 'level'.
    """

    tags:list[dict] = []
    #sample tags strcture:
    # tags = [{'start':0, 'end':3, 'structure':'opening', 'type':'p', 'attributes':(r'type="blablabla"'), 'level':0}]

    for index, char in enumerate(text):
        # Extracts information about each tag such as structure, type, attributes, etc.


        #gets opening and closing index (index of < & >)
        if char != '<':
            continue

        closing_index = text.find('>', index + 1)
        
    #gets structure variable
        structure:str | None = None
        if text[index + 1] == r'/': #</tag>
            structure = 'closing'
        elif text[closing_index - 1] == r'/': #<tag/>
            structure = 'self-closing'
        else: #<tag>
            structure = 'opening'

        #get the content
        #first get range of the content in between <>
        #Then using that range gets a string of that content

        content_range:tuple[int,int] = ()  #refers to the content in between the <>
        if structure == 'opening':
            content_range = (index + 1, closing_index)  
        else:
            content_range = (index + 2, closing_index)
            if structure == 'self-closing':
                content_range = (index + 1, closing_index - 1)

        content:str | None = None
        start , end = content_range
        _content:list[str] = []

        for i, v in enumerate(range(start,end)):
            _content.append(text[v])

        content = ''.join(_content)

# get the tag type variable <tag>
        #                    ^^^
        content_separated:list[str] = content.split()
        type = content_separated[0]
        attributes:list[str] = []

#puts attributes into a list of strs
        if len(content_separated) > 1:
            for i,attribute in enumerate(content_separated[1:]):
                attributes.append(attribute)

        tag_dictionary:dict[str,any] = {'start':index, 'end':closing_index, 'structure':structure, 'type':type, 'attributes':attributes}
        tags.append(tag_dictionary)

    level: int = 0
    prev_structure: str|None = None

    for tag in tags:
        """
        ALGORITHM:

        1. Starts at level 0.
        2. If tag structure is same as old structure, level increases by 1.
        3. Else if it is different, level decreases by 1.
        4. If tag structure is self-closing, it treats its structure as if it is the same as the old structure, But returns self-closing to old structure variable.
        5. Finally, if a self-closing structured tag, precedes a closing tag, level decreases by 1. 
        """

        tag_structure = tag.get('structure')
        old_level = level

        def update_level(level:int) -> int:
            if prev_structure == 'closing':
                level -= 1
            if prev_structure == 'opening':
                level += 1
            
            return level

        if tag_structure == 'self-closing':
            # level = update_level(level)
            tag['level'] = level
            # prev_structure = tag_structure

            continue
    
        level = update_level(level)

        if prev_structure == 'self-closing':
            if tag_structure == 'closing':
                level -= 1
            if tag_structure == 'opening':
                level += 1

            tag['level'] = level
            prev_structure = tag_structure

            continue

        if tag_structure != prev_structure :
            level = old_level

        tag['level'] = level
        prev_structure = tag_structure

        continue

    return tags

def pair_tags(tags:list[dict[str, any]]) ->list[tuple[dict[str, any], dict[str, any]|None]]:
    """
    Pair opening and closing tags from a list of tags.

    This function pairs opening tags by finding the next closing tag of same type and level.

    Parameters:
    tags (list): A list of dictionaries representing tags with keys 'structure', 'type', 'level', and 'start'.

    Returns:
    list: A list of tuples containing paired tags where the second element is None for self-closing tags.
    """
    paired_tags:list[tuple[dict,dict]] = []

    _tags:list[dict[str, any]] = tags[:]  # Create a copy of the tags list

    for tag in _tags:
        tag_structure = tag.get('structure')
        tag_type = tag.get('type')
        tag_level = tag.get('level')
        tag_start = tag.get('start')

        if tag_structure == 'self-closing':
            paired_tags.append((tag, None))
            continue

        if tag_structure == 'closing':
            continue

        # Find the next tag with the same type and level
        condition = lambda tag: (
            tag.get('type') == tag_type and
            tag.get('level') == tag_level and
            tag.get('structure') == 'closing' and
            tag.get('start') > tag_start
        )

        closing_tag = next((_tag for _tag in _tags if condition(_tag)), None)

        paired_tags.append((tag, closing_tag))

    pass
    return paired_tags

def get_elements(tags:list[tuple[dict,dict]], text:str) -> list[dict[str, tuple[int, int]]]:
    """
    Extracts elements from the given text and their corresponding parent tags.

    Parameters:
    tags (list): A list of dictionaries representing paired tags with keys 'structure', 'type', 'level', and 'start'.
    text (str): The text content from which to extract elements.

    Returns:
    list: A list of dictionaries representing elements with keys 'element', 'element_range', and 'parent_tags'.

    This function processes the text content to extract individual elements and determine their parent tags based on the provided tag pairs. It splits the text into elements using BeautifulSoup, identifies the index ranges for each tag pair, and then extracts the elements and their index ranges. Finally, it assigns the corresponding parent tags to each element based on the tag pair ranges.

    The output is a list of dictionaries, each representing a element with the following keys:
    - 'element': The extracted element content.
    - 'element_range': A tuple containing the start and end indices of the element within the text.
    - 'parent_tags': A list of IDs representing the parent tags associated with the element.
    """

    # Create a new list to hold the combined elements and self-closing tags
    arranged_elements:list[str] = []

    # This block fills in the 'arranged_elements' list
    #made just to avoid repitition
    def node_append(text_body:str)->None:
        bs4_elements_to_search:BeautifulSoup = BeautifulSoup(text_body, 'html.parser')
        text_nodes:list = bs4_elements_to_search.find_all(text=True)
        text_nodes = [text_node.replace('\n','').strip() for text_node in text_nodes if text_node.replace('\n','').strip()]
        arranged_elements.extend(text_nodes)

    search_start:int = 0
    for tag in tags: #finds the nearest self closing tag and uses its location in the xml to  get all the text before it. 
        tag1 = tag[0]
        if tag1.get('structure') !='self-closing':
            continue

        self_closing_tags_string:str = f'<{tag1.get('type')} {' '.join(tag1.get('attributes'))}/>' 

        search_end = tag1.get('start')
        text_to_search = text[search_start:search_end]
        search_start = search_end

        node_append(text_to_search)
        arranged_elements.append(self_closing_tags_string)

    last_string_search:str = text[search_start:]
    node_append(last_string_search)

    # List to store the index ranges for each tag pair
    tag_pair_ranges: list[dict[str, id | tuple[int, int]]] = []

    # Iterate through tag pairs to fill in the tag_pair_ranges dictionary
    for tag_pair in tags:
        range_end: int = None
        range_start: int = None
        _tag_pair_container: dict[str, id | tuple[int, int]] = {}

        if tag_pair[0].get('structure') == 'self-closing':
            continue

        # Determine the start and end index for the tag pair
        if tag_pair[0] and isinstance(tag_pair[0], dict):
            range_start = tag_pair[0].get('end')
        if tag_pair[1] and isinstance(tag_pair[1], dict):
            range_end = tag_pair[1].get('start')

        # Store the index range in a tuple
        _range: tuple[int, int] = (range_start, range_end)
        _tag_pair_container['properties'] = tag_pair
        _tag_pair_container['tag_pair_range'] = _range
        tag_pair_ranges.append(_tag_pair_container)

    # Extracts the elements and their index ranges
    elements_ranges: list[dict[str, tuple[int,int]]] = []
    elements_to_skip:list[str] = []
    _text:str = str(text) 

    for element in sorted(arranged_elements,key=len, reverse=True):
        if element in elements_to_skip:
            continue
        element_range:dict[str,str|tuple[int,int]] = {}
        occs:int = _text.count(element)
        prev_index:int = 0
        current_index:int = 0
        for _ in range(occs):
            current_index = _text.find(element, current_index + prev_index)
            end_index:int = current_index + len(element)
            _text = list(_text)

            for i in range(current_index, end_index): # do not turn to one liner
                _text[i] = '*'

            _text = ''.join(_text)
            element_range:dict = {'element': element, 'range': (current_index, end_index)}
            elements_ranges.append(element_range)
            prev_index = current_index

        elements_to_skip.append(element)

    elements_ranges = sorted(elements_ranges, key=lambda element_range: element_range.get('range')[0])

    # Add the parent tags for each element range
    for element_range in elements_ranges:
         
        parent_tags: list[id] = []
        element_range_start, element_range_end = element_range.get('range')
        element = element_range.get('element')
        for pair_range in tag_pair_ranges:
             
            tag_range_start, tag_range_end = pair_range.get('tag_pair_range')
            tag_pair_properties = pair_range.get('properties')
             
            if element_range_start > tag_range_start and element_range_end < tag_range_end:
                parent_tags.append(tag_pair_properties)
        element_range['parent_tags'] = parent_tags 

    return elements_ranges 
    # AFTER THIS FUNCTION IS EXECUTED ALL THAT NEEDS TO BE DONE IS FOR EACH ELEMENT TO BE INSTANTIATED AS A

def docElementinstantiator(elements,tags,document) -> None:
    #PHASE 1:
    #first separate each element into groups based on its first paragraph parent and if its a list
    # make sure that elements maintain chronological order
    #NOTE never have nested paragraph or a list in p tag or a list in a list tag

    paragraphTags = [tag for tag in tags if tag[0].get('type') in ('p', 'list')] #puts all paragraph tags in a list

    paragraph_groups = []
    list_groups = []
    for paragraphTag in paragraphTags:
        pTag_range = (paragraphTag[0].get('start'), paragraphTag[1].get('end'))
        isList = paragraphTag[0].get('type') == 'list'
        same_parent_paragraph = []
        for i, element in enumerate(elements):
            isLastiteration = i == len(elements) - 1
            element_range = element.get('range')
            if pTag_range[0] < element_range[0] < pTag_range[1] and pTag_range[1] > element_range[1] > pTag_range[0]: #Basically, if range (w,x) is within (y,z)
                same_parent_paragraph.append(element)
            
            if isLastiteration and isList:
                list_groups.append(same_parent_paragraph)
            elif isLastiteration and not isList:
                paragraph_groups.append(same_parent_paragraph)
                

    ...
    #PHASE 2:
    #Once a list has been created for this, instatiate a paragraph for each list then keep that in a dict like so:
    #{paragraph: paragraph_elemnt, elements: [...]}
    #store these dicts in a list
    instantiatedParagraphs = []
    for paragraphGroup in paragraph_groups:
        docParagraph = document.add_paragraph() #FOR FUTURE PURPOSES: if need be add paragraph styles here
        instantiatedParagraphs.append({'paragraph': docParagraph, 'elements': paragraphGroup, 'list': False})
    for listGroup in list_groups:
        docParagraph = document.add_paragraph()
        instantiatedParagraphs.append({'paragraph': docParagraph, 'elements': listGroup, 'list': True})
    ...

    #PHASE 3:
    #In this phase the actual elements are being created in the file itself
    #assuming that the code for the actual creation is already made (done in a class)
    #The only purpose of this phase is to ensure that they are made in the correct order\
    
    def create_docList(paragrpah, elements, list_style='bullet'):
        ORDERED:list[str] = ['num', 'alpha', '_alpha', 'roman', '_roman'] 
        #num: numbered 123456  alpha: ABCDEF _alpha: abcdef roman: I II III IV _roman :i ii iii vi 
        UNORDERED:list['str'] = ['dash', 'bullet', 'hollow_bullet']
        # dash: -  bullet: • hollow_bullet: ο

        if list_style in ORDERED:
            ...
        elif list_style in UNORDERED:
            ...
        else:
            raise ValueError (f'List style is {list_style} expected {zip(ORDERED,UNORDERED)}')
        

    for paragraphInstance in instantiatedParagraphs:
        paragraph = paragraphInstance.get('paragraph')
        child_elements = paragraphInstance.get('elements')
        if paragraphInstance.get('list'):
            ... #TODO: create a list class where when given parent paragraph(that acts like a list) and child elements 
            # with given parent elements it creates a list ex. docList(paragraphInstance, elements, list style)
        elif not paragraphInstance.get('list'):
            for element in child_elements:
                #Adds the runs to each paragraph
                run:DocRun = DocRun(paragraph,element) #instantiation
                run.addtoRootParagraph()

def create_document(xml):
    doc = Document()

    current_time = time.localtime()
    formatted_time = time.strftime(r"%Y%m%dT%H%M%S", current_time)

    #requests to add in the document
    tags:list[tuple[dict]] = get_tags(xml)

    tags = pair_tags(tags)

    elements:list[dict] = get_elements(tags,xml)

    docElementinstantiator(elements, tags, doc)

    save_directory:str = get_target_dir()

    last_folder:str = os.path.basename(os.path.normpath(save_directory))
    optional_name:str = input('Optional name: ')
    file_name = f"{optional_name}-{last_folder}-{formatted_time}.docx"

    doc.save(fr'{save_directory}/{file_name}')



def main():
    with open('documentxml.txt', 'r') as xml:
        xml_body = xml.read()
        create_document(xml_body)


if __name__ == '__main__':
    main()